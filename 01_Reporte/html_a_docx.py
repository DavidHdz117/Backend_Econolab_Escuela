from __future__ import annotations

import argparse
import re
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


RED = "A80000"
DARK = "202124"
LIGHT_RED = "FFF4F4"
LIGHT_GRAY = "F3F3F3"


def shade(element, color: str) -> None:
    properties = element.get_or_add_tcPr() if hasattr(element, "get_or_add_tcPr") else element.get_or_add_pPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def prevent_row_split(row) -> None:
    row_properties = row._tr.get_or_add_trPr()
    if row_properties.find(qn("w:cantSplit")) is None:
        row_properties.append(OxmlElement("w:cantSplit"))


def repeat_header(row) -> None:
    row_properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    row_properties.append(header)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("Página ")
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, end))


class HtmlReportConverter:
    def __init__(self, html_path: Path, output_path: Path) -> None:
        self.html_path = html_path.resolve()
        self.output_path = output_path.resolve()
        self.document = Document()
        self._configure_document()

    def _configure_document(self) -> None:
        section = self.document.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

        normal = self.document.styles["Normal"]
        normal.font.name = "Aptos"
        normal.font.size = Pt(10.5)
        normal.font.color.rgb = RGBColor.from_string(DARK)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.15

        heading_sizes = {1: 20, 2: 15, 3: 12.5, 4: 11}
        for level, size in heading_sizes.items():
            style = self.document.styles[f"Heading {level}"]
            style.font.name = "Aptos Display"
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = RGBColor.from_string(RED if level <= 2 else DARK)
            style.paragraph_format.keep_with_next = True
            style.paragraph_format.space_before = Pt(14 if level <= 2 else 9)
            style.paragraph_format.space_after = Pt(5)

        if "Code Block" not in [style.name for style in self.document.styles]:
            style = self.document.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = "Consolas"
            style.font.size = Pt(8)
            style.font.color.rgb = RGBColor.from_string("FFFFFF")
            style.paragraph_format.space_before = Pt(5)
            style.paragraph_format.space_after = Pt(7)

        header = section.header.paragraphs[0]
        header.text = "ECONOLAB · Proyecto práctico de Machine Learning"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string("666666")
        add_page_field(section.footer.paragraphs[0])

        properties = self.document.core_properties
        properties.title = "Reporte práctico de Machine Learning — ECONOLAB"
        properties.subject = "Extracción de conocimiento en bases de datos"
        properties.author = "Equipo ECONOLAB"

    @staticmethod
    def _classes(tag: Tag) -> set[str]:
        return set(tag.get("class", []))

    @staticmethod
    def _clean_text(value: str) -> str:
        return re.sub(r"[ \t\r\f\v]+", " ", value).strip()

    def _add_inline(
        self,
        paragraph,
        node,
        *,
        bold: bool = False,
        italic: bool = False,
        code: bool = False,
    ) -> None:
        if isinstance(node, NavigableString):
            text = re.sub(r"\s+", " ", str(node))
            if text:
                run = paragraph.add_run(text)
                run.bold = bold
                run.italic = italic
                if code:
                    run.font.name = "Consolas"
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = RGBColor.from_string(RED)
            return
        if not isinstance(node, Tag):
            return
        if node.name == "br":
            paragraph.add_run().add_break()
            return
        next_bold = bold or node.name in {"strong", "b"}
        next_italic = italic or node.name in {"em", "i"}
        next_code = code or node.name == "code"
        for child in node.children:
            self._add_inline(
                paragraph,
                child,
                bold=next_bold,
                italic=next_italic,
                code=next_code,
            )

    def _add_paragraph(self, tag: Tag, *, style: str | None = None):
        paragraph = self.document.add_paragraph(style=style)
        self._add_inline(paragraph, tag)
        classes = self._classes(tag)
        if "center" in classes:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif tag.name == "p":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if "small" in classes:
            for run in paragraph.runs:
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor.from_string("666666")
        return paragraph

    def _add_heading(self, tag: Tag) -> None:
        if "page-break" in self._classes(tag):
            self.document.add_page_break()
        level = int(tag.name[1])
        paragraph = self.document.add_heading(level=level)
        self._add_inline(paragraph, tag)

    def _add_table(self, tag: Tag) -> None:
        html_rows = tag.find_all("tr")
        if not html_rows:
            return
        column_count = max(
            len(row.find_all(["th", "td"], recursive=False)) for row in html_rows
        )
        if column_count == 0:
            return
        table = self.document.add_table(rows=len(html_rows), cols=column_count)
        table.style = "Table Grid"
        table.autofit = True
        for row_index, html_row in enumerate(html_rows):
            cells = html_row.find_all(["th", "td"], recursive=False)
            row = table.rows[row_index]
            prevent_row_split(row)
            if row_index == 0:
                repeat_header(row)
            for column_index, html_cell in enumerate(cells):
                cell = row.cells[column_index]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                paragraph = cell.paragraphs[0]
                paragraph.paragraph_format.space_after = Pt(1.5)
                self._add_inline(paragraph, html_cell)
                is_header = html_cell.name == "th"
                if is_header:
                    shade(cell._tc, RED)
                for run in paragraph.runs:
                    run.font.name = "Aptos"
                    run.font.size = Pt(8)
                    if is_header:
                        run.bold = True
                        run.font.color.rgb = RGBColor.from_string("FFFFFF")
        self.document.add_paragraph().paragraph_format.space_after = Pt(1)

    def _resolve_image(self, source: str) -> Path:
        return (self.html_path.parent / source).resolve()

    def _add_figure(self, tag: Tag) -> None:
        image = tag.find("img")
        if image and image.get("src"):
            path = self._resolve_image(image["src"])
            if path.is_file():
                paragraph = self.document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                width = Inches(3.2) if "econolab-brand" in path.name else Inches(6.25)
                run.add_picture(str(path), width=width)
        caption = tag.find("figcaption")
        if caption:
            paragraph = self._add_paragraph(caption)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.italic = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor.from_string("666666")

    def _add_callout(self, tag: Tag, color: str) -> None:
        paragraph = self._add_paragraph(tag)
        shade(paragraph._p, color)
        paragraph.paragraph_format.left_indent = Cm(0.25)
        paragraph.paragraph_format.right_indent = Cm(0.25)
        paragraph.paragraph_format.space_before = Pt(5)
        paragraph.paragraph_format.space_after = Pt(7)

    def _add_pre(self, tag: Tag) -> None:
        paragraph = self.document.add_paragraph(style="Code Block")
        paragraph.add_run(tag.get_text()).font.name = "Consolas"
        shade(paragraph._p, "111111")

    def _add_list(self, tag: Tag) -> None:
        style = "List Number" if tag.name == "ol" else "List Bullet"
        for item in tag.find_all("li", recursive=False):
            paragraph = self.document.add_paragraph(style=style)
            self._add_inline(paragraph, item)

    def _add_cover(self, tag: Tag) -> None:
        for child in tag.children:
            if not isinstance(child, Tag):
                continue
            if child.name == "img" and child.get("src"):
                path = self._resolve_image(child["src"])
                if path.is_file():
                    paragraph = self.document.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.add_run().add_picture(str(path), width=Inches(3.2))
            elif child.name in {"h1", "h2"}:
                paragraph = self.document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._add_inline(paragraph, child, bold=True)
                for run in paragraph.runs:
                    run.font.name = "Aptos Display"
                    run.font.size = Pt(25 if child.name == "h1" else 17)
                    run.font.color.rgb = RGBColor.from_string(RED if child.name == "h1" else DARK)
                paragraph.paragraph_format.space_after = Pt(12)
            elif child.name == "p":
                paragraph = self._add_paragraph(child)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif child.name == "table":
                self._add_table(child)
        self.document.add_page_break()

    def _render_block(self, tag: Tag) -> None:
        if tag.name == "section" and "cover" in self._classes(tag):
            self._add_cover(tag)
        elif tag.name in {"h1", "h2", "h3", "h4"}:
            self._add_heading(tag)
        elif tag.name == "p":
            self._add_paragraph(tag)
        elif tag.name in {"ol", "ul"}:
            self._add_list(tag)
        elif tag.name == "table":
            self._add_table(tag)
        elif tag.name == "figure":
            self._add_figure(tag)
        elif tag.name == "pre":
            self._add_pre(tag)
        elif tag.name == "div":
            classes = self._classes(tag)
            if "callout" in classes:
                self._add_callout(tag, LIGHT_GRAY)
            elif "warning" in classes or "pending" in classes:
                self._add_callout(tag, LIGHT_RED)
            elif "success" in classes:
                self._add_callout(tag, "F3F8F3")
            else:
                for child in tag.children:
                    if isinstance(child, Tag):
                        self._render_block(child)

    def convert(self) -> None:
        html = self.html_path.read_text(encoding="utf-8")
        soup = BeautifulSoup(html, "html.parser")
        if soup.body is None:
            raise ValueError("El HTML no contiene body.")
        for child in soup.body.children:
            if isinstance(child, Tag):
                self._render_block(child)

        final_section = self.document.sections[-1]
        if final_section.start_type == WD_SECTION.NEW_PAGE:
            final_section.start_type = WD_SECTION.CONTINUOUS
        self.output_path.parent.mkdir(parents=True, exist_ok=True)
        self.document.save(self.output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("html", type=Path)
    parser.add_argument("docx", type=Path)
    args = parser.parse_args()
    HtmlReportConverter(args.html, args.docx).convert()
    print(f"DOCX generado: {args.docx.resolve()}")


if __name__ == "__main__":
    main()
